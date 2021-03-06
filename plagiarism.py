

import xlrd
import docx
import os
import time
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Inches
import xlsxwriter
import xlwt
import shutil
from os import chdir, getcwd, listdir, path
from datetime import datetime
# import sounddevice as sd
from tkinter import *
from PIL import ImageTk, Image
from tkinter import filedialog
from tkinter.filedialog import askdirectory
from tkinter import messagebox
import sys
import numpy as np
from math import pi
import requests
import client
import tkinter as tk
from tkinter import messagebox

# from final import firstfile , secondfile
# if(fileloc = true)
# open log file, ******* process statr
# else
# * create a log file open  ***** process starte

start_time = time.time()
file_location = " "

search_dir = " "
loc = " "
first_name = []
roll_no = []
student_name = []
assingnment_code = []
first_file = []
first_len = 0
space = []
student_id_col = 0
student_name_col = 0
assignment_code_col = 0
document_link_col = 0
selected = ""
rejected = ""
unchecked = ""

bad_chars = [';', ':', '!', "*", '-', ',', '?', '+', '-', '*', '/']


def create_folder(report):
    global selected
    global rejected
    global unchecked
    selected = report + "/1.Selected Assignment"
    rejected = report + "/2.Rejected Assignment"
    unchecked = report + "/3.Un Checked Assignment"

    if not os.path.exists(selected):
        os.mkdir(selected)

    if not os.path.exists(rejected):
        os.mkdir(rejected)

    if not os.path.exists(unchecked):
        os.mkdir(unchecked)


def column_to_number(c):
    number = 0
    for l in c:
        number = (ord(l.upper()) - 64) + (number * 26)
    return number - 1


def Excel_format():
    master = tk.Tk()

    tk.Label(master, text="Enter Student ID column number").grid(row=0)
    tk.Label(master, text="Enter Student name column number").grid(row=2)
    tk.Label(master, text="Enter Assignment code column number").grid(row=4)
    tk.Label(master, text="Enter document link column number").grid(row=6)

    e1 = tk.Entry(master)
    e2 = tk.Entry(master)
    e3 = tk.Entry(master)
    e4 = tk.Entry(master)

    def submit():
        global student_id_col
        student_id_col = column_to_number(e1.get())

        global student_name_col
        student_name_col = column_to_number(e2.get())

        global assignment_code_col
        assignment_code_col = column_to_number(e3.get())

        global document_link_col
        document_link_col = column_to_number(e4.get())

        print("student_id_col: %d\n student_name_col : %d\nassignment_code_col : %d\ndocument_link_col : %d\n" % (
        student_id_col, student_name_col, assignment_code_col, document_link_col))
        master.destroy()

    e1.grid(row=0, column=1, padx=10, pady=10)
    e2.grid(row=2, column=1, padx=10, pady=10)
    e3.grid(row=4, column=1, padx=10, pady=10)
    e4.grid(row=6, column=1, padx=10, pady=10)

    tk.Button(master, text='Submit', command=submit).grid(row=10, column=1, sticky=tk.W, pady=4)
    tk.mainloop()


def Excel_File():
    global file_location
    file_location = filedialog.askopenfilename()
    if (not (file_location.endswith('.xlsx'))):
        messagebox.showwarning("Warning", "Choose Excel File Only")
        file_path = " "
    else:
        # global workbook
        # workbook=xlrd.open_workbook(file_location)
        print(file_location)


def Folder_upload():
    global search_dir
    search_dir = askdirectory()
    print(search_dir)


def firstfile(location):
    doc = docx.Document(location)
    global first_file
    first_file = []
    a = " "
    for i in doc.paragraphs:
        # a.append(i.text)
        a += " " + i.text
        a = ''.join(j for j in a if not j in bad_chars)
    list = a.split()
    global first_len
    first_len = len(list)
    # print(first_len)

    for i in range(0, first_len):
        if ((first_len - i) > 2):
            x = list[i] + " " + list[i + 1] + " " + list[i + 2]
            first_file.append(x)
        # print(x)


# print("\n first file \n\n")
# print(first_file)

def secondfile(location):
    second_file = []
    doc = docx.Document(location)
    b = " "
    for i in doc.paragraphs:
        # a.append(i.text)
        b += " " + i.text
        b = ''.join(j for j in b if not j in bad_chars)

    list = b.split()
    second_len = len(list)
    for i in range(0, second_len):
        if ((second_len - i) > 2):
            x = list[i] + " " + list[i + 1] + " " + list[i + 2]
            second_file.append(x)
    # print("\n second file \n\n")
    # print(second_file)
    similarity_file = ""
    global first_len
    second_len = len(second_file)
    # print("first_file",len(first_file))
    # print("second_file",len(second_file))
    i = 0
    first_len = len(first_file)
    for n in range(0, len(first_file)):
        # for j in range(0,len(second_file)):
        if (first_len > i and first_file[i] in second_file):
            # print("File 1",first_file[i])
            similarity_file += " " + first_file[i]
            i += 2
        i = i + 1
    # print(similarity_file)
    temp = similarity_file.split()
    # print("temp= ",len(temp))
    # print("first len",first_len)
    try:
        similarity = (len(temp) / first_len) * 100
        similarity = float("{0:.2f}".format(round(similarity, 2)))
    except ZeroDivisionError:
        similarity = 0
    # print(similarity)
    return similarity


def check():
    create_folder(loc)
    print(selected)
    print(rejected)
    print(unchecked)
    print(file_location)
    print(search_dir)
    print(loc)

    if (file_location == " " and search_dir == " "):
        messagebox.showwarning("Warning", "Choose Excel Sheet and Folder Upload")
    elif (file_location == " "):
        messagebox.showwarning("Warning", "Choose Excel Sheet")
    elif (search_dir == " "):
        messagebox.showwarning("Warning", "Choose Folder Upload")
    elif (loc == " "):
        messagebox.showwarning("Warning", "Choose Location to save  plagiarism report")
    else:
        workbook1 = xlsxwriter.Workbook(loc + "/Result.xlsx")
        worksheet = workbook1.add_worksheet('Result')
        worksheet.set_column('A:A', 20)
        worksheet.write("A1", 'PDF Link')
        worksheet.write("B1", 'Selected/Rejected')
        worksheet.write("C1", 'Max Similarity')

        workbook = xlrd.open_workbook(file_location)
        sheet = workbook.sheet_by_index(0)
        no_of_rows = sheet.nrows
        print(student_id_col, student_name_col, assignment_code_col, document_link_col)
        for i in range(1, no_of_rows):
            # print(len(sheet.cell_value(i,13)))

            if (len(sheet.cell_value(i, document_link_col)) != 0):
                x = (sheet.cell_value(i, student_id_col))
                y = (sheet.cell_value(i, student_name_col))
                z = (sheet.cell_value(i, assignment_code_col))
                # print(i,x)
                roll_no.append(x)
                student_name.append(y)
                assingnment_code.append(z)
            else:
                space.append(i)

        print(space)

        # search_dir = "C:/Users/Logesh kc/Downloads/DOC file upload (File responses)-20190917T091312Z-001/DOC file upload (File responses)"
        os.chdir(search_dir)
        files = filter(os.path.isfile, os.listdir(search_dir))
        files = [os.path.join(search_dir, f) for f in files]
        files.sort(key=lambda x: os.path.getmtime(x))
        print(files)
        j = 1

        flag = 0
        print(len(files))
        for i in range(0, len(files)):
            records = []
            max = 0.0
            start_time = time.time()
            print("\n    ", i, "   ")
            print(files[i])
            if (not files[i].endswith(".docx")) or (files[i].startswith("~$")):
                shutil.copy(files[i], unchecked)
                continue

            firstfile(files[i])

            if (i == 0):
                similarity = 0

            for j in range(i - 1, -1, -1):

                list = []
                if (not files[j].endswith(".docx")) or (files[i].startswith("~$")):
                    continue
                similarity = secondfile(files[j])
                if (similarity >= 100.0):
                    similarity = 100.0

                list.append(roll_no[j])
                list.append(student_name[j])
                list.append(assingnment_code[j])
                list.append(similarity)

                if (similarity > max):
                    max = similarity
                records.append(list)

            document = Document()
            document.add_picture('C:/Plagiarism/clg.png', width=Inches(5.00))
            # document.add_picture('clg.png')
            document.add_heading('  \tBIT Assignment Plagiarism Report', 0).bold = True

            if (max > 30.0):
                result = 'Rejected'
                reason = 'Similarity percentage >  30'
            else:
                result = "Selected"
                reason = "None"

            table = document.add_table(rows=7, cols=2)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Student ID'
            hdr_cells[1].text = ':  ' + str(roll_no[i])

            hdr_cells = table.rows[1].cells
            hdr_cells[0].text = 'Student Name'
            hdr_cells[1].text = ':  ' + str(student_name[i])

            hdr_cells = table.rows[2].cells
            hdr_cells[0].text = 'Assignment'
            hdr_cells[1].text = ':  ' + str(assingnment_code[i])

            hdr_cells = table.rows[3].cells
            hdr_cells[0].text = 'Maximum Similarity'
            hdr_cells[1].text = ':  ' + str(max)

            hdr_cells = table.rows[4].cells
            hdr_cells[0].text = 'Result (Selected \ Rejected)'
            hdr_cells[1].text = ':  ' + result

            hdr_cells = table.rows[5].cells
            hdr_cells[0].text = 'Reason'
            hdr_cells[1].text = ':  ' + reason

            hdr_cells = table.rows[6].cells
            hdr_cells[0].text = 'Report Generated on'
            hdr_cells[1].text = ':  ' + str(datetime.now().strftime('%Y-%m-%d    %H:%M:%S'))

            print("Maximum Similarity = ", max)

            if (i != 0):

                table = document.add_table(rows=1, cols=5, style='Table Grid')
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'S.No'
                hdr_cells[1].text = 'Student Id'
                hdr_cells[2].text = 'Student Name'
                hdr_cells[3].text = 'Assignment'
                hdr_cells[4].text = '% Similarity'
                count = 1

                for qty, id, desc, simi in records:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(count)
                    row_cells[1].text = str(qty)
                    row_cells[2].text = str(id)
                    row_cells[3].text = str(desc)
                    row_cells[4].text = str(simi)
                    count += 1

            #foo = document.add_paragraph().add_run(
             #   '''Report Generated by Open Source Special Lab\nFaculty Guide Ms. Malathi T(AP-CSE)\nDeveloped by Logesh kc & Karthickkumar K (CSE-2017-21)''')
            #font = foo.font
            #font.color.rgb = RGBColor(255, 0, 0)

            # doc=loc+"/demo.docx"

            # document.save(doc)

            # file = open(doc, 'rb')
            # data = file.read()
            # res = requests.post(url='http://converter-eval.plutext.com:80/v1/00000000-0000-0000-0000-000000000000/convert',
            #          data=data,
            #          headers={'Content-Type': 'application/octet-stream'})

            # pdf=loc+"/"+str(i)+"."+roll_no[i]+".pdf"
            # f = open(pdf, 'wb')
            # f.write(res.content)
            # f.close()

            if ((i + 1 + flag) in space):
                # print("i",i,"\nflag",flag)
                flag += 1
            x = "A" + str(i + flag + 2)
            y = "B" + str(i + flag + 2)
            z = "C" + str(i + flag + 2)
            # print("i",i,"\nflag",flag,"\n x",x)
            # worksheet.write_url(x,pdf,string=roll_no[i])
            if (result == "Rejected"):
                shutil.copy(files[i], rejected)
                # pdf=rejected+"/"+str(i)+"."+roll_no[i]+".pdf"
                doc = rejected + "/." + str(i) + "." + str(roll_no[i]) + ".docx"
                document.save(doc)
                worksheet.write_url(x, doc, string=str(roll_no[i]))
                # f = open(pdf, 'wb')
                # f.write(res.content)
                # f.close()

                worksheet.write(y, result)
            else:
                shutil.copy(files[i], selected)
                # pdf=selected+"/"+str(i)+"."+roll_no[i]+".pdf"
                doc = selected + "/." + str(i) + "." + str(roll_no[i]) + ".docx"
                document.save(doc)
                worksheet.write_url(x, doc, string=str(roll_no[i]))
                # f= open(pdf, 'wb')
                # f.write(res.content)
                # f.close()

                worksheet.write(y, result)

            worksheet.write(z, max)
            print("End")
            end_time = time.time()
            print(i, " = Execution time=", end_time - start_time)

        workbook1.close()
        # document.close()
        # os.remove(doc)
        end_time = time.time()
        print("Execution time=", end_time - start_time)

        messagebox.showinfo(" ", "Report Generated Completed")
        root.destroy()

def checkentry():
    if (file_location == " " and search_dir == " "):
        messagebox.showwarning("Warning", "Choose Excel Sheet and Folder Upload")
    else:
        os.chdir(search_dir)
        files = filter(os.path.isfile, os.listdir(search_dir))
        files = [os.path.join(search_dir, f) for f in files]
        files.sort(key=lambda x: os.path.getmtime(x))
        print(files)

        workbook = xlrd.open_workbook(file_location)
        sheet = workbook.sheet_by_index(0)
        no_of_rows = sheet.nrows

        if(len(files)==no_of_rows-1):
            messagebox.showinfo(" ","Given Entries are matched Successfully..!!")
        else:
            messagebox.showinfo(" ","The Entries between Excel file and Uploaded documents are mismatched, CAN'T GENERATE PLAGIARISM REPORT")

        root.destroy()


def Exit():
    root.destroy()


def loc():
    global loc
    loc = askdirectory()
    print(loc)


root = Tk()
root.title("BIT Internal Assignment Plagiarism Identifier (BIAPI)")
root.geometry("550x400")
root.resizable(0, 0)

#menubar = Menu(root)
#filemenu = Menu(menubar, tearoff=0)
#filemenu.add_command(label="New", )

#menubar.add_cascade(label="File", menu=filemenu)

#helpmenu = Menu(menubar, tearoff=0)
#helpmenu.add_command(label="Help ", )
#helpmenu.add_command(label="About...", )
#menubar.add_cascade(label="Help", menu=helpmenu)

#root.config(menu=menubar)

left = Frame(root, relief="solid", bg="gray")
left.pack(side="left", expand=True, fill="both")
right = Frame(root, relief="solid", )
left.pack(side="right", expand=True, fill="both")

# right.pack(fill= "both")
container = Frame(left, relief="solid")
box01 = Frame(right, relief="solid")
box1 = Frame(right, relief="solid")
box2 = Frame(right, relief="solid")
box3 = Frame(right, relief="solid")
box4 = Frame(right, relief="solid")
box5 = Frame(right, relief="solid")
box6 = Frame(right, relief="solid")

canvas = Canvas(container, width=250, height=265)
canvas.pack()
img = ImageTk.PhotoImage(Image.open("bg2.png"))
canvas.create_image(0, 0, anchor=NW, image=img)

Button01 = Button(box01, text="Excel Format", width=30, command=Excel_format, foreground="black", background="white")
Button1 = Button(box1, text="Location to save  plagiarism report", width=30, command=loc, foreground="black",
                 background="white")
Button2 = Button(box2, text="Upload Excel Sheet", width=30, command=Excel_File, foreground="black", background="white")
Button3 = Button(box3, text="Upload Assignment Folder ", width=30, command=Folder_upload, foreground="black",
                 background="white")
Button4 = Button(box4, text="Generate Plagiarism Report", width=30, command=check, foreground="black",
                 background="white")
Button5 = Button(box5, text="Exit", width=30, command=Exit, foreground="black", background="white")
Button6 = Button(box6, text="CHECK", width=30, command=checkentry, foreground="black",
                 background="white")
# Button(root, text = 'Click Me !', image = PhotoImage,).pack()

left.pack(side="left", expand=True, fill="both")
right.pack(side="right", expand=True)
container.pack(expand=True)

box01.pack(expand=True, padx=10, pady=10)
box2.pack(expand=True, padx=10, pady=10)
box3.pack(expand=True, padx=10, pady=10)
box6.pack(expand=True, padx=10, pady=10)
box1.pack(expand=True, padx=10, pady=10)
box4.pack(expand=True, padx=10, pady=10)
box5.pack(expand=True, padx=10, pady=10)

Button01.pack()
Button1.pack()
Button2.pack()
Button3.pack()
Button4.pack()
Button5.pack()
Button6.pack()

status = Label(left, text="Developed in Open Source Lab", bd=1, relief="solid", anchor=W)
status.pack(side=BOTTOM, fill=X)

root.mainloop()